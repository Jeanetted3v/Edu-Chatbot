from typing import Dict, Union, List
import os
import logging
import pypdf
import pandas as pd
from omegaconf import DictConfig
from docx import Document
from dataclasses import dataclass


logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class LoadedUnstructuredDocument:
    """Represents a loaded document with its content and metadata."""
    content: str
    metadata: Dict[str, str]


@dataclass
class LoadedStructuredDocument:
    """Represents a loaded structured document (e.g., CSV, Excel) with its content and metadata."""
    content: pd.DataFrame
    metadata: Dict[str, str]


class LocalDocLoader:
    """Loads documents of various formats (PDF, DOCX, TXT) into a unified format."""
    @staticmethod
    def convert_excel_to_csv(excel_path: str) -> str:
        """
        Convert Excel file to CSV and save in the same directory
        
        Args:
            excel_path: Path to the Excel file
            
        Returns:
            Path to the saved CSV file
        """
        try:
            # Get directory and filename without extension
            directory = os.path.dirname(excel_path)
            base_name = os.path.splitext(os.path.basename(excel_path))[0]
            csv_path = os.path.join(directory, f"{base_name}.csv")
            
            # Read Excel and save as CSV
            df = pd.read_excel(excel_path)
            df.to_csv(csv_path, index=False, encoding='utf-8-sig')
            
            logger.info(f"Successfully converted Excel to CSV: {csv_path}")
            return csv_path
            
        except Exception as e:
            logger.error(f"Error converting Excel to CSV: {str(e)}")
            raise
    
    def _load_csv(self, file_path: str) -> LoadedStructuredDocument:
        """
        Load a CSV file into a pandas DataFrame.
        Attempts to handle common CSV format variations.
        """
        metadata = {
            'source': file_path,
            'type': 'csv'
        }

        # Try different encodings
        encodings = ['utf-8-sig', 'utf-8', 'latin1', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                df = pd.read_csv(file_path, encoding=encoding)
                
                metadata.update({
                    'encoding': encoding,
                    'rows': str(len(df)),
                    'columns': str(len(df.columns)),
                    'column_names': ', '.join(df.columns.tolist())
                })
                
                return LoadedStructuredDocument(content=df, metadata=metadata)
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not load CSV file: {file_path}. Tried encodings: {encodings}")

    def _load_pdf(self, file_path: str) -> LoadedUnstructuredDocument:
        """Load a PDF document."""
        metadata = {
            'source': file_path,
            'type': 'pdf'
        }
        try: 
            with open(file_path, 'rb') as file:
                pdf = pypdf.PdfReader(file)
                if len(pdf.pages) == 0:
                    raise ValueError(f"PDF file {file_path} is empty")
                content = []
                total_chars = 0
                
                for i, page in enumerate(pdf.pages, 1):
                    try: 
                        text = page.extract_text()
                        if text:
                            content.append(text)
                            chars_in_page = len(text)
                            total_chars += chars_in_page
                            metadata[f'page_{i}_length'] = str(chars_in_page)
                
                    except Exception as e:
                        logger.warning(f"Error reading page {i} of PDF {file_path}: {str(e)}")
                metadata['total_pages'] = str(len(pdf.pages))
                
                full_text = "\n\n".join(content)
                
                if not full_text.strip():
                    raise ValueError(f"No text could be extracted from PDF {file_path}")
                
                return LoadedUnstructuredDocument(content=full_text, metadata=metadata)       
        except FileNotFoundError:
            raise FileNotFoundError(f"PDF file not found: {file_path}")
        except Exception as e:
            raise ValueError(f"Error reading PDF {file_path}: {str(e)}")
    
    def _load_docx(self, file_path: str) -> LoadedUnstructuredDocument:
        """Load a DOCX document."""
        metadata = {
            'source': file_path,
            'type': 'docx'
        }
        
        doc = Document(file_path)
        paragraphs = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        metadata['total_paragraphs'] = str(len(paragraphs))
        
        return LoadedUnstructuredDocument(content=full_text, metadata=metadata)
    
    def _load_text(self, file_path: str) -> LoadedUnstructuredDocument:
        """Load a text document."""
        metadata = {
            'source': file_path,
            'type': 'txt'
        }
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        metadata['file_size'] = str(os.path.getsize(file_path))
        
        return LoadedUnstructuredDocument(content=content, metadata=metadata)

    def _load_document(self, file_path: str) -> Union[
        LoadedUnstructuredDocument,
        LoadedStructuredDocument
    ]:
        """
        Load a document and return its contents with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            LoadedDocument object containing text content and metadata
        
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._load_pdf(file_path)
        elif file_ext == '.docx':
            return self._load_docx(file_path)
        elif file_ext == '.txt':
            return self._load_text(file_path)
        # Structured document handlers
        elif file_ext == '.csv':
            return self._load_csv(file_path)
        elif file_ext in ['.xlsx', '.xls']:
            csv_path = self.convert_excel_to_csv(file_path)
            return self._load_csv(csv_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")


def load_local_doc(
    cfg: DictConfig
) -> List[Union[LoadedUnstructuredDocument, LoadedStructuredDocument]]:
    """
    Load documents from local filesystem based on configuration.
    
    Args:
        cfg: Hydra configuration object
        doc_loader: Instance of DocLoader
        
    Returns:
        List of loaded documents
    """
    documents = []
    
    # Load documents from configured paths
    for doc_path in cfg.LOCAL_DOC.PATHS:
        try:
            doc_loader = LocalDocLoader()
            doc = doc_loader.load_document(doc_path)
            documents.append(doc)
            logger.info(f"Successfully loaded document: {doc_path}")
            
            # Log preview and metadata
            if isinstance(doc, LoadedUnstructuredDocument):
                logger.debug(f"Content preview: {doc.content[:200]}...")
                logger.debug(f"Metadata: {doc.metadata}")
            
        except Exception as e:
            logger.error(f"Error loading document {doc_path}: {str(e)}")
            
    return documents
